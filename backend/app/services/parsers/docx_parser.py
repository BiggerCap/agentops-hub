"""
DOCX Parser - Extract text from Word documents
"""

from docx import Document


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text content from a DOCX file
    
    Args:
        file_path: Absolute path to the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        Exception: If DOCX reading fails
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise Exception("No text could be extracted from DOCX")
        
        return full_text
        
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
